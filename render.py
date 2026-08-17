"""Turn Markdown resumes/letters into files an ATS will actually accept.

DOCX is the safest universal upload format for ATS parsers. HTML is provided so
you can print to PDF from a browser without extra dependencies.
"""
from __future__ import annotations

import html
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITAL = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def safe_slug(*parts: str) -> str:
    s = "_".join(p.strip() for p in parts if p)
    s = re.sub(r"[^\w\-. ]+", "", s).strip().replace(" ", "_")
    return re.sub(r"_{2,}", "_", s)[:80] or "application"


# --------------------------------------------------------------------- DOCX

def _add_runs(paragraph, text: str) -> None:
    """Minimal inline markdown -> docx runs (bold/italic/links-as-text)."""
    text = _LINK.sub(r"\1 (\2)", text)
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    rest = text[pos:]
    pos2 = 0
    for m in _ITAL.finditer(rest):
        if m.start() > pos2:
            paragraph.add_run(rest[pos2:m.start()])
        paragraph.add_run(m.group(1)).italic = True
        pos2 = m.end()
    if rest[pos2:]:
        paragraph.add_run(rest[pos2:])


def markdown_to_docx(md: str, out_path: str | Path, base_font: str = "Calibri") -> Path:
    if Document is None:
        raise RuntimeError("python-docx not installed: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = base_font
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(50)

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(11)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip().upper())
            r.bold = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(19)
        elif line.lstrip().startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, line.lstrip()[2:].strip())
        elif set(line.strip()) <= {"-", "_", "="} and len(line.strip()) >= 3:
            continue  # horizontal rule
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_runs(p, line.strip())

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --------------------------------------------------------------------- HTML

_CSS = """
body{font-family:-apple-system,Segoe UI,Helvetica,Arial,sans-serif;max-width:760px;
margin:40px auto;padding:0 24px;color:#111;line-height:1.45;font-size:14px}
h1{text-align:center;margin-bottom:4px;font-size:26px}
h2{font-size:13px;letter-spacing:.08em;text-transform:uppercase;color:#1f3a5f;
border-bottom:1px solid #d5d5d5;padding-bottom:3px;margin-top:22px}
h3{font-size:14px;margin:12px 0 2px}
ul{margin:4px 0 8px 18px;padding:0}li{margin:2px 0}
@media print{body{margin:0;font-size:11.5px}@page{margin:14mm}}
"""


def markdown_to_html(md: str, out_path: str | Path, title: str = "Resume") -> Path:
    lines, out, in_list = md.splitlines(), [], False

    def inline(t: str) -> str:
        t = html.escape(t)
        t = _LINK.sub(r'<a href="\2">\1</a>', t)
        t = _BOLD.sub(r"<strong>\1</strong>", t)
        return _ITAL.sub(r"<em>\1</em>", t)

    for raw in lines:
        line = raw.rstrip()
        is_li = line.lstrip().startswith(("- ", "* ", "• "))
        if in_list and not is_li:
            out.append("</ul>")
            in_list = False
        if not line.strip():
            continue
        if is_li:
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{inline(line.lstrip()[2:].strip())}</li>")
        elif line.startswith("### "):
            out.append(f"<h3>{inline(line[4:])}</h3>")
        elif line.startswith("## "):
            out.append(f"<h2>{inline(line[3:])}</h2>")
        elif line.startswith("# "):
            out.append(f"<h1>{inline(line[2:])}</h1>")
        else:
            out.append(f"<p>{inline(line)}</p>")
    if in_list:
        out.append("</ul>")

    doc = (f"<!doctype html><meta charset='utf-8'><title>{html.escape(title)}</title>"
           f"<style>{_CSS}</style>" + "\n".join(out))
    p = Path(out_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(doc, encoding="utf-8")
    return p


def write_package(app, out_dir: str | Path) -> dict[str, str]:
    """Write resume + letter in md/html/docx. Returns kind -> path."""
    folder = Path(out_dir) / safe_slug(app.company, app.title, app.job_id[:6])
    folder.mkdir(parents=True, exist_ok=True)
    artifacts: dict[str, str] = {}

    stem = safe_slug(app.autofill.get("full_name", "candidate"), app.company)

    (folder / f"{stem}_resume.md").write_text(app.resume_md, encoding="utf-8")
    artifacts["resume_md"] = str(folder / f"{stem}_resume.md")
    artifacts["resume_html"] = str(
        markdown_to_html(app.resume_md, folder / f"{stem}_resume.html", "Resume"))

    (folder / f"{stem}_cover_letter.md").write_text(app.cover_letter_md, encoding="utf-8")
    artifacts["cover_md"] = str(folder / f"{stem}_cover_letter.md")

    if Document is not None:
        try:
            artifacts["resume_docx"] = str(
                markdown_to_docx(app.resume_md, folder / f"{stem}_resume.docx"))
            artifacts["cover_docx"] = str(
                markdown_to_docx(app.cover_letter_md, folder / f"{stem}_cover_letter.docx"))
        except Exception as e:  # noqa: BLE001
            artifacts["docx_error"] = str(e)

    return artifacts
